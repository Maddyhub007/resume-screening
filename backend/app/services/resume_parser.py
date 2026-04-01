"""
app/services/resume_parser.py  ── IMPROVED VERSION

Multi-format resume parser.

Supported formats:  PDF (pdfplumber primary, PyMuPDF fallback),  DOCX.

Pipeline:
  1. Extract raw text from file.
  2. Segment text into resume sections (skills, experience, education, etc.).
  3. Extract structured data from each section.
  4. Compute derived metrics (total_experience_years, skill_count).
  5. Return a ParseResult dataclass.

IMPROVEMENTS OVER ORIGINAL
---------------------------
  Phase 1 — High-impact fixes (no new dependencies):
    [C1]  DOCX: doc.tables now extracted (tabular skills/experience recovered).
    [C2]  DOCX: headers & footers extracted (contact info no longer lost).
    [C3]  _SECTION_PATTERNS expanded from 6 → 50+ variants (WORK HISTORY,
          KEY SKILLS, INTERNSHIPS, etc. now matched).
    [C4]  Skill vocabulary imported from skill_taxonomy.py (400+ entries).
    [C5]  C++ / C# boundary regex fixed (non-word-char lookaround).
    [C6]  Full month names ("January", "February") parsed in date ranges.
    [C7]  MM/YYYY date format ("06/2022 – Present") now handled.
    [C8]  B.Tech / B.E. / B.Com / PG Diploma added to education regex.
    [C9]  Negation filter (8-token window) prevents false positives.

  Phase 2 — PDF layout improvements:
    [C10] pdfplumber column detection: x-coordinate grouping reassembles
          two-column PDFs in correct reading order.
    [C11] pdfplumber table extraction: skills in PDF tables now recovered.
    [C12] 'to' word date separator now matched ("June 2020 to March 2022").

  Phase 3 — Enrichment:
    [C13] OOV (out-of-vocabulary) skill capture from comma/bullet lists.
    [C14] Contact info extraction (_extract_contact).
    [C15] Hyphenated skill forms matched ("scikit-learn", "test-driven").

Design decisions:
  - Two-library PDF strategy: pdfplumber is more accurate for text-heavy PDFs;
    PyMuPDF (fitz) handles scanned-ish or complex layout PDFs better.  We try
    pdfplumber first and fall back to fitz on empty output.
  - Section detection uses keyword-based heuristics — no ML required.  Fast,
    deterministic, and works offline.
  - Skill extraction: exact-match + normalisation against skill_taxonomy vocab.
  - Experience parser: extracts date ranges from each job entry and sums months.
  - All parsing errors are caught and surfaced in parse_error so the caller can
    set parse_status=FAILED on the Resume record without crashing.

Usage:
    from app.services.resume_parser import ResumeParserService
    svc = ResumeParserService()
    result = svc.parse("/path/to/resume.pdf")
    result = svc.parse_bytes(file_bytes, filename="resume.docx")
"""

import logging
import os
import re
import tempfile
from dataclasses import dataclass, field
from typing import Optional

logger = logging.getLogger(__name__)

# Import expanded taxonomy (falls back gracefully if file not yet deployed)
try:
    from app.services.skill_taxonomy import (
        TECH_SKILLS as _TECH_SKILLS,
        SKILL_ALIASES as _SKILL_ALIASES,
        SPECIAL_BOUNDARY_SKILLS as _SPECIAL_BOUNDARY_SKILLS,
        NEGATION_WORDS as _NEGATION_WORDS,
    )
except ImportError:
    # Fallback to original minimal vocabulary if taxonomy module missing
    logger.warning("skill_taxonomy.py not found — using minimal built-in vocab")
    _TECH_SKILLS: set[str] = {
        "python", "javascript", "typescript", "java", "c++", "c#",
        "react", "vue", "angular", "node.js", "flask", "django",
        "tensorflow", "pytorch", "pandas", "sql", "postgresql",
        "aws", "gcp", "azure", "docker", "kubernetes", "git",
    }
    _SKILL_ALIASES: dict[str, str] = {
        "js": "javascript", "ts": "typescript", "py": "python",
        "node": "node.js", "k8s": "kubernetes", "ml": "machine learning",
    }
    _SPECIAL_BOUNDARY_SKILLS: dict[str, str] = {
        "c++": r"(?<![a-zA-Z0-9])c\+\+(?![a-zA-Z0-9])",
        "c#":  r"(?<![a-zA-Z0-9])c#(?![a-zA-Z0-9])",
    }
    _NEGATION_WORDS: frozenset[str] = frozenset({
        "no", "not", "without", "except", "excluding", "lacks",
    })


# ─────────────────────────────────────────────────────────────────────────────
# Result dataclass  (NEW: contact field added)
# ─────────────────────────────────────────────────────────────────────────────

@dataclass
class ParseResult:
    """Structured output of a resume parse operation."""

    raw_text: str = ""
    skills: list[str] = field(default_factory=list)
    education: list[dict] = field(default_factory=list)
    experience: list[dict] = field(default_factory=list)
    certifications: list[str] = field(default_factory=list)
    projects: list[dict] = field(default_factory=list)
    summary_text: str = ""
    total_experience_years: float = 0.0
    skill_count: int = 0
    contact: dict = field(default_factory=dict)   # [C14] NEW
    oov_skills: list[str] = field(default_factory=list)  # [C13] NEW
    parse_error: Optional[str] = None

    @property
    def success(self) -> bool:
        return self.parse_error is None and bool(self.raw_text)


# ─────────────────────────────────────────────────────────────────────────────
# Section header patterns  [C3] EXPANDED from 6 → 50+ variants
# ─────────────────────────────────────────────────────────────────────────────

_SECTION_PATTERNS = {
    "summary": re.compile(
        r"^\s*(summary|professional\s+summary|career\s+summary|profile|"
        r"professional\s+profile|career\s+profile|objective|career\s+objective|"
        r"about\s+me|personal\s+statement|executive\s+summary|overview|"
        r"professional\s+overview|introduction)\s*$",
        re.I,
    ),
    "skills": re.compile(
        r"^\s*(skills?|technical\s+skills?|core\s+competencies?|"
        r"key\s+skills?|professional\s+skills?|technologies?|"
        r"tools?\s*(&|and)?\s*technologies?|tech\s+stack|"
        r"technical\s+expertise|areas?\s+of\s+expertise|"
        r"expertise|competencies?|programming\s+languages?|"
        r"languages?\s*(&|and)?\s*frameworks?|"
        r"frameworks?\s*(&|and)?\s*libraries?|"
        r"software\s+skills?|hard\s+skills?)\s*$",
        re.I,
    ),
    "experience": re.compile(
        r"^\s*(experience|work\s+experience|professional\s+experience|"
        r"employment|employment\s+history|work\s+history|career\s+history|"
        r"professional\s+history|job\s+history|"
        r"internship|internships?|work\s+placements?|placements?|"
        r"industry\s+experience|relevant\s+experience|"
        r"positions?\s+held|roles?\s+held)\s*$",
        re.I,
    ),
    "education": re.compile(
        r"^\s*(education|educational\s+background|academic\s+background|"
        r"academic\s+qualifications?|qualifications?|degrees?|schooling|"
        r"academics?|university|college|institutes?|"
        r"training\s*(&|and)?\s*education)\s*$",
        re.I,
    ),
    "certifications": re.compile(
        r"^\s*(certifications?|certificates?|licenses?|accreditations?|"
        r"professional\s+certifications?|industry\s+certifications?|"
        r"credentials?|awarded\s+credentials?|badges?)\s*$",
        re.I,
    ),
    "projects": re.compile(
        r"^\s*(projects?|personal\s+projects?|open\s+source|portfolio|"
        r"side\s+projects?|key\s+projects?|notable\s+projects?|"
        r"academic\s+projects?|course\s+projects?|"
        r"capstone|thesis|dissertations?)\s*$",
        re.I,
    ),
    "achievements": re.compile(
        r"^\s*(achievements?|accomplishments?|awards?|honours?|honors?|"
        r"recognitions?|milestones?)\s*$",
        re.I,
    ),
    "languages": re.compile(
        r"^\s*(languages?|spoken\s+languages?|linguistic\s+skills?)\s*$",
        re.I,
    ),
}


# ─────────────────────────────────────────────────────────────────────────────
# Date patterns  [C6] full month names  [C7] MM/YYYY  [C12] 'to' separator
# ─────────────────────────────────────────────────────────────────────────────

# Accepts abbreviated AND full month names ("Jan", "January")
_DATE_RANGE_RE = re.compile(
    r"(jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|"
    r"jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|"
    r"nov(?:ember)?|dec(?:ember)?)"
    r"[\s,\.]+(\d{4})"
    r"\s*[-–—to/]+\s*"
    r"(jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|"
    r"jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|"
    r"nov(?:ember)?|dec(?:ember)?|present|current|now)"
    r"(?:[\s,\.]+(\d{4}))?",
    re.I,
)

# [C7] MM/YYYY – MM/YYYY  or  MM/YYYY – Present
_MM_YYYY_RANGE_RE = re.compile(
    r"(\d{1,2})/(\d{4})\s*[-–—to]+\s*(?:(\d{1,2})/(\d{4})|(present|current|now))",
    re.I,
)

_YEAR_RANGE_RE = re.compile(
    r"(\d{4})\s*[-–—to]+\s*(\d{4}|present|current|now)", re.I
)

_MONTH_MAP = {
    "jan": 1, "january": 1,
    "feb": 2, "february": 2,
    "mar": 3, "march": 3,
    "apr": 4, "april": 4,
    "may": 5,
    "jun": 6, "june": 6,
    "jul": 7, "july": 7,
    "aug": 8, "august": 8,
    "sep": 9, "september": 9,
    "oct": 10, "october": 10,
    "nov": 11, "november": 11,
    "dec": 12, "december": 12,
}

# ─────────────────────────────────────────────────────────────────────────────
# Contact patterns  [C14]
# ─────────────────────────────────────────────────────────────────────────────

_EMAIL_RE    = re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}")
_PHONE_RE    = re.compile(
    r"(?:\+?\d{1,3}[\s\-.]?)?"
    r"(?:\(?\d{2,4}\)?[\s\-.]?)"
    r"\d{3,4}[\s\-.]?\d{3,4}"
)
_LINKEDIN_RE = re.compile(r"linkedin\.com/in/[\w\-]+", re.I)
_GITHUB_RE   = re.compile(r"github\.com/[\w\-]+", re.I)


# ─────────────────────────────────────────────────────────────────────────────
# Education degree patterns  [C8] B.Tech, B.E., B.Com, PG Diploma added
# ─────────────────────────────────────────────────────────────────────────────

_DEGREE_RE = re.compile(
    r"(bachelor|master|phd|ph\.d|doctorate|associate|diploma|"
    r"pg\s+diploma|post\s+graduate\s+diploma|postgraduate\s+diploma|"
    r"b\.?tech|b\.?e\.?|b\.?com|b\.?sc?|m\.?sc?|m\.?tech|m\.?e\.?|"
    r"m\.?eng?|b\.?eng?|mba|b\.?a\.?|m\.?a\.?|"
    r"bachelor\s+of\s+\w+|master\s+of\s+\w+)",
    re.I,
)


# ─────────────────────────────────────────────────────────────────────────────
# Pre-compile skill search patterns for performance
# ─────────────────────────────────────────────────────────────────────────────

def _build_skill_patterns() -> list[tuple[str, re.Pattern]]:
    """
    Build compiled regex patterns for every skill in the vocabulary.

    Skills with special characters (c++, c#, etc.) use custom lookaround
    patterns [C5].  All others use standard \b word boundaries.
    """
    patterns = []
    # Sort longest-first so multi-word skills match before substrings
    for skill in sorted(_TECH_SKILLS, key=len, reverse=True):
        if skill in _SPECIAL_BOUNDARY_SKILLS:
            pat = re.compile(_SPECIAL_BOUNDARY_SKILLS[skill], re.I)
        else:
            escaped = re.escape(skill)
            pat = re.compile(r"\b" + escaped + r"\b", re.I)
        patterns.append((skill, pat))
    return patterns


_SKILL_PATTERNS: list[tuple[str, re.Pattern]] = _build_skill_patterns()


# ─────────────────────────────────────────────────────────────────────────────
# Main service
# ─────────────────────────────────────────────────────────────────────────────

class ResumeParserService:
    """
    Multi-format resume parser.

    Stateless — safe to share across threads / requests.
    """

    def parse(self, file_path: str) -> ParseResult:
        """
        Parse a resume from a file path.

        Args:
            file_path: Absolute or relative path to a .pdf or .docx file.

        Returns:
            ParseResult — always returns, never raises.
        """
        if not os.path.exists(file_path):
            return ParseResult(parse_error=f"File not found: {file_path}")

        ext = os.path.splitext(file_path)[1].lower()
        try:
            if ext == ".pdf":
                raw_text = self._extract_pdf(file_path)
            elif ext == ".docx":
                raw_text = self._extract_docx(file_path)
            else:
                return ParseResult(parse_error=f"Unsupported file type: {ext}")
        except Exception as exc:
            logger.exception("Text extraction failed for %s", file_path)
            return ParseResult(parse_error=f"Extraction error: {exc}")

        if not raw_text or not raw_text.strip():
            return ParseResult(
                raw_text=raw_text or "",
                parse_error="No text could be extracted from the file.",
            )

        return self._parse_text(raw_text)

    def parse_bytes(self, data: bytes, filename: str) -> ParseResult:
        """
        Parse a resume from raw bytes (e.g., uploaded file data).

        Writes to a temp file, parses, then cleans up.
        """
        ext = os.path.splitext(filename)[1].lower()
        tmp_path = None
        try:
            with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
                tmp.write(data)
                tmp_path = tmp.name
            result = self.parse(tmp_path)
        finally:
            if tmp_path:
                try:
                    os.unlink(tmp_path)
                except Exception:
                    pass
        return result

    # ── PDF extraction ────────────────────────────────────────────────────────

    def _extract_pdf(self, path: str) -> str:
        """Try pdfplumber (with column detection); fall back to PyMuPDF."""
        text = self._extract_pdf_pdfplumber(path)
        if text and len(text.strip()) > 100:
            return text
        logger.debug("pdfplumber yielded sparse text, trying PyMuPDF fallback.")
        return self._extract_pdf_fitz(path)

    @staticmethod
    def _extract_pdf_pdfplumber(path: str) -> str:
        """
        Extract text from PDF using pdfplumber.

        Improvements [C10][C11]:
          - Column detection: words are grouped by x-position into left/right
            columns and reassembled in reading order.
          - Table extraction: page.extract_tables() is now called and text
            from table cells is appended to the page text.
        """
        try:
            import pdfplumber
            pages_text = []
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    # [C11] Extract table cells first
                    table_text = ""
                    try:
                        tables = page.extract_tables()
                        if tables:
                            cell_texts = []
                            for table in tables:
                                for row in table:
                                    for cell in row:
                                        if cell and str(cell).strip():
                                            cell_texts.append(str(cell).strip())
                            if cell_texts:
                                table_text = "\n".join(cell_texts)
                    except Exception:
                        pass

                    # [C10] Column-aware text extraction
                    page_text = ResumeParserService._extract_page_with_columns(page)
                    if not page_text:
                        page_text = page.extract_text() or ""

                    combined = (page_text.strip() + "\n" + table_text.strip()).strip()
                    if combined:
                        pages_text.append(combined)

            return "\n".join(pages_text)
        except ImportError:
            logger.debug("pdfplumber not available.")
            return ""
        except Exception as exc:
            logger.debug("pdfplumber extraction error: %s", exc)
            return ""

    @staticmethod
    def _extract_page_with_columns(page) -> str:
        """
        [C10] Detect two-column layout by x-coordinate clustering.

        If the page words cluster into two x-groups (left/right columns),
        extract each column separately in top-to-bottom order, then join them.
        This prevents interleaving like "Sr Engineer Data Scientist" instead
        of separate left-column and right-column entries.
        """
        try:
            words = page.extract_words()
            if not words:
                return ""

            # Find the median x0 to split columns
            x0_values = sorted(w["x0"] for w in words)
            if not x0_values:
                return page.extract_text() or ""

            page_width = page.width or 600
            mid = page_width / 2

            # Count words on each side; if skewed, it's single column
            left_count  = sum(1 for x in x0_values if x < mid)
            right_count = len(x0_values) - left_count

            # Consider two-column if both sides have ≥20% of words
            ratio = min(left_count, right_count) / max(len(x0_values), 1)
            if ratio < 0.2:
                # Single column — standard extraction
                return page.extract_text() or ""

            # Two-column: split at page midpoint, sort each by top
            left_words  = [w for w in words if w["x0"] < mid]
            right_words = [w for w in words if w["x0"] >= mid]

            def words_to_text(word_list):
                if not word_list:
                    return ""
                # Sort by top (y position), then x
                sorted_words = sorted(word_list, key=lambda w: (round(w["top"], 1), w["x0"]))
                lines = []
                current_line = []
                current_top = None
                for w in sorted_words:
                    top = round(w["top"], 1)
                    if current_top is None or abs(top - current_top) <= 3:
                        current_line.append(w["text"])
                        current_top = top
                    else:
                        if current_line:
                            lines.append(" ".join(current_line))
                        current_line = [w["text"]]
                        current_top = top
                if current_line:
                    lines.append(" ".join(current_line))
                return "\n".join(lines)

            left_text  = words_to_text(left_words)
            right_text = words_to_text(right_words)
            return left_text + "\n" + right_text

        except Exception as exc:
            logger.debug("Column extraction failed: %s", exc)
            return ""

    @staticmethod
    def _extract_pdf_fitz(path: str) -> str:
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(path)
            pages = [doc[i].get_text() for i in range(len(doc))]
            doc.close()
            return "\n".join(pages)
        except ImportError:
            logger.debug("PyMuPDF (fitz) not available.")
            return ""
        except Exception as exc:
            logger.debug("PyMuPDF extraction error: %s", exc)
            return ""

    # ── DOCX extraction  [C1][C2] ─────────────────────────────────────────────

    @staticmethod
    def _extract_docx(path: str) -> str:
        """
        [C1] Extract text from DOCX including tables.
        [C2] Extract text from headers and footers (contact info lives there).

        Original only iterated doc.paragraphs — tables and headers/footers
        were completely skipped.
        """
        try:
            from docx import Document
            doc = Document(path)
            parts: list[str] = []

            # [C2] Headers and footers from every section
            for section in doc.sections:
                for hdr_or_ftr in (section.header, section.footer):
                    try:
                        for para in hdr_or_ftr.paragraphs:
                            t = para.text.strip()
                            if t:
                                parts.append(t)
                    except Exception:
                        pass

            # Body paragraphs (original behaviour)
            for para in doc.paragraphs:
                t = para.text.strip()
                if t:
                    parts.append(t)

            # [C1] Tables — iterate every cell
            for table in doc.tables:
                for row in table.rows:
                    row_cells = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_cells.append(cell_text)
                    if row_cells:
                        parts.append("  |  ".join(row_cells))

            return "\n".join(parts)
        except ImportError:
            return ""
        except Exception as exc:
            logger.debug("DOCX extraction error: %s", exc)
            return ""

    # ── Text parsing pipeline ─────────────────────────────────────────────────

    def _parse_text(self, raw_text: str) -> ParseResult:
        """Run the full NLP pipeline on extracted text."""
        sections = self._segment_sections(raw_text)

        # Search skills section + first 500 chars of full text (header area)
        skills_text = sections.get("skills", "") + "\n" + raw_text[:500]
        skills, oov_skills = self._extract_skills(skills_text)

        education       = self._extract_education(sections.get("education", ""))
        experience      = self._extract_experience(sections.get("experience", ""))
        certifications  = self._extract_certifications(sections.get("certifications", ""))
        projects        = self._extract_projects(sections.get("projects", ""))
        summary         = sections.get("summary", "")
        contact         = self._extract_contact(raw_text[:1000])  # [C14]

        total_exp_years = self._compute_experience_years(experience)

        return ParseResult(
            raw_text=raw_text,
            skills=skills,
            education=education,
            experience=experience,
            certifications=certifications,
            projects=projects,
            summary_text=summary,
            total_experience_years=round(total_exp_years, 1),
            skill_count=len(skills),
            contact=contact,
            oov_skills=oov_skills,
        )

    # ── Section segmentation ──────────────────────────────────────────────────

    def _segment_sections(self, text: str) -> dict[str, str]:
        """
        Split resume text into labelled sections.

        Returns dict mapping section_name → section_text.
        """
        lines = text.split("\n")
        sections: dict[str, str] = {}
        current_section = "header"
        buffer: list[str] = []

        for line in lines:
            stripped = line.strip()
            matched_section = None
            for name, pattern in _SECTION_PATTERNS.items():
                if pattern.match(stripped):
                    matched_section = name
                    break

            if matched_section:
                sections[current_section] = "\n".join(buffer).strip()
                current_section = matched_section
                buffer = []
            else:
                buffer.append(line)

        sections[current_section] = "\n".join(buffer).strip()
        return sections

    # ── Skills extraction  [C4][C5][C9][C13][C15] ────────────────────────────

    def _extract_skills(self, text: str) -> tuple[list[str], list[str]]:
        """
        Extract skills from the skills section + header.

        Returns (known_skills, oov_skills) where:
          - known_skills: matched against TECH_SKILLS vocabulary, negation-filtered
          - oov_skills:   comma/bullet tokens NOT in vocabulary (candidates for
                          future additions) [C13]

        Improvements:
          [C4]  Uses expanded 400+ entry vocabulary from skill_taxonomy.py.
          [C5]  Special boundary regex for c++, c#, .net etc.
          [C9]  Negation filter: skips skill if preceded by negation word.
          [C13] OOV skill capture from comma/bullet separated lists.
          [C15] Hyphenated forms match via \b boundary.
        """
        text_lower = text.lower()
        found: set[str] = set()

        # Split text into tokens for negation window check
        tokens = text_lower.split()

        for skill, pattern in _SKILL_PATTERNS:
            for match in pattern.finditer(text_lower):
                # [C9] Negation check: look at up to 8 tokens before match
                match_pos = match.start()
                prefix = text_lower[:match_pos]
                prefix_tokens = prefix.split()[-8:]
                if any(neg in prefix_tokens for neg in _NEGATION_WORDS):
                    continue
                found.add(skill)
                break  # skill found, no need to check further matches

        # Alias normalisation
        normalised: set[str] = set()
        for s in found:
            normalised.add(_SKILL_ALIASES.get(s, s))

        # [C13] OOV skill capture from comma/bullet-delimited lines
        oov_skills = self._capture_oov_skills(text, normalised)

        return sorted(normalised), sorted(oov_skills)

    @staticmethod
    def _capture_oov_skills(text: str, known_skills: set[str]) -> list[str]:
        """
        [C13] Capture unknown skills from comma or bullet-separated lines.

        Heuristic: if a line has 3+ comma-separated tokens (or bullet items)
        and the tokens are short (likely skill names), extract tokens not
        already in the known vocabulary.
        """
        oov: set[str] = set()
        stop_words = {
            "and", "or", "the", "a", "an", "with", "for", "in", "of",
            "to", "on", "at", "by", "from", "including", "etc",
        }

        for line in text.split("\n"):
            line = line.strip().lstrip("•-*·▪►").strip()
            if not line:
                continue

            # Try comma-separated
            candidates = [t.strip() for t in re.split(r"[,;|]", line) if t.strip()]
            if len(candidates) >= 3:
                for candidate in candidates:
                    candidate_lower = candidate.lower().strip("()[].")
                    if (
                        2 <= len(candidate_lower) <= 40
                        and candidate_lower not in stop_words
                        and candidate_lower not in known_skills
                        and not re.match(r"^\d+$", candidate_lower)
                    ):
                        oov.add(candidate_lower)

        return sorted(oov)[:20]  # cap at 20 to avoid noise

    # ── Contact extraction  [C14] ─────────────────────────────────────────────

    @staticmethod
    def _extract_contact(text: str) -> dict:
        """
        [C14] Extract contact information from the top of the resume.

        Looks for email, phone, LinkedIn URL, and GitHub URL.
        """
        contact: dict = {}

        email_match = _EMAIL_RE.search(text)
        if email_match:
            contact["email"] = email_match.group(0)

        phone_match = _PHONE_RE.search(text)
        if phone_match:
            # Basic sanity check: at least 7 digits
            digits = re.sub(r"\D", "", phone_match.group(0))
            if len(digits) >= 7:
                contact["phone"] = phone_match.group(0).strip()

        linkedin_match = _LINKEDIN_RE.search(text)
        if linkedin_match:
            contact["linkedin"] = "https://www." + linkedin_match.group(0)

        github_match = _GITHUB_RE.search(text)
        if github_match:
            contact["github"] = "https://" + github_match.group(0)

        return contact

    # ── Education extraction  [C8] ────────────────────────────────────────────

    def _extract_education(self, text: str) -> list[dict]:
        """
        [C8] Extract education entries — now recognises B.Tech, B.E., B.Com,
        PG Diploma in addition to original patterns.
        """
        entries = []
        if not text.strip():
            return entries

        year_re = re.compile(r"\b(19|20)\d{2}\b")
        lines = [l.strip() for l in text.split("\n") if l.strip()]

        for i, line in enumerate(lines):
            if _DEGREE_RE.search(line):
                years = year_re.findall(line)
                institution = ""
                if i + 1 < len(lines):
                    institution = lines[i + 1]
                entries.append({
                    "degree":      line[:200],
                    "institution": institution[:200],
                    "year":        years[-1] if years else "",
                })
                if len(entries) >= 5:
                    break

        return entries

    # ── Experience extraction ─────────────────────────────────────────────────

    def _extract_experience(self, text: str) -> list[dict]:
        """Extract job experience entries."""
        entries = []
        if not text.strip():
            return entries

        # Try double-newline split first (well-formatted resumes)
        blocks = re.split(r"\n{2,}", text.strip())

        # If only one block returned, the resume is single-spaced
        # Fall back to splitting on lines that look like job titles
        if len(blocks) <= 1:
            blocks = self._split_experience_by_title(text)

        for block in blocks[:10]:
            block = block.strip()
            if len(block) < 20:
                continue
            lines = [l.strip() for l in block.split("\n") if l.strip()]
            if not lines:
                continue

            title_line   = lines[0]
            company_line = lines[1] if len(lines) > 1 else ""
            date_line    = ""
            description_lines = []

            date_re = re.compile(r"(19|20)\d{2}|present|current", re.I)
            for j, ln in enumerate(lines[1:], start=1):
                if date_re.search(ln):
                    date_line = ln
                    description_lines = lines[j + 1:]
                    break
            else:
                description_lines = lines[2:]

            entries.append({
                "title":       title_line[:200],
                "company":     company_line[:200],
                "date_range":  date_line[:100],
                "description": " ".join(description_lines)[:500],
            })

        return entries

    @staticmethod
    def _split_experience_by_title(text: str) -> list[str]:
        """
        Split single-spaced experience section by detecting title-like lines.
        A title line is typically short, title-cased, and followed by a company.
        """
        lines = text.split("\n")
        blocks = []
        current: list[str] = []

        # Heuristic: a new entry starts when we see a short ALL-CAPS or
        # Title Case line that is NOT a bullet point
        title_re = re.compile(r"^[A-Z][a-zA-Z\s,\-/&]+$")
        date_re  = re.compile(r"(19|20)\d{2}|present|current", re.I)

        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue
            is_title_like = (
                title_re.match(stripped)
                and len(stripped) < 80
                and not stripped.startswith(("•", "-", "*", "·"))
                and not date_re.search(stripped)
            )
            if is_title_like and current:
                blocks.append("\n".join(current))
                current = [stripped]
            else:
                current.append(stripped)

        if current:
            blocks.append("\n".join(current))

        return blocks if blocks else [text]

    # ── Certifications extraction ─────────────────────────────────────────────

    def _extract_certifications(self, text: str) -> list[str]:
        """Extract certification names from the certifications section."""
        if not text.strip():
            return []
        certs = []
        for line in text.split("\n"):
            line = line.strip().lstrip("•-*·").strip()
            if len(line) > 5:
                certs.append(line[:200])
        return certs[:10]

    # ── Projects extraction ───────────────────────────────────────────────────

    def _extract_projects(self, text: str) -> list[dict]:
        """Extract project entries from the projects section."""
        if not text.strip():
            return []
        projects = []
        blocks = re.split(r"\n{2,}", text.strip())
        for block in blocks[:8]:
            lines = [l.strip() for l in block.split("\n") if l.strip()]
            if not lines:
                continue
            projects.append({
                "name":        lines[0][:200],
                "description": " ".join(lines[1:])[:400],
            })
        return projects

    # ── Experience years computation  [C6][C7][C12] ───────────────────────────

    def _compute_experience_years(self, experience_entries: list[dict]) -> float:
        """
        Sum up total months of experience across all job entries.

        Uses fuzzy date parsing:
          1. Month+year ranges ("Jan 2020 – Present")  [C6] full month names
          2. MM/YYYY ranges ("06/2022 – Present")      [C7]
          3. Year-only ranges ("2020 – 2023")
        """
        import datetime
        now = datetime.date.today()
        total_months = 0

        for entry in experience_entries:
            date_str = entry.get("date_range", "")
            months = self._parse_duration_months(date_str, now)
            total_months += months

        return min(total_months / 12.0, 40.0)  # cap at 40 years

    def _parse_duration_months(self, date_str: str, now) -> int:
        """Parse a date range string and return duration in months."""
        import datetime
        if not date_str:
            return 0

        # [C6] Try full/abbreviated month+year range
        m = _DATE_RANGE_RE.search(date_str)
        if m:
            try:
                start_key = m.group(1).lower()
                start_month = _MONTH_MAP.get(start_key, 1)
                start_year  = int(m.group(2))
                end_token   = m.group(3).lower()
                if end_token.startswith(("pre", "cur", "now")):
                    end_date = now
                else:
                    end_month = _MONTH_MAP.get(end_token, 1)
                    end_year  = int(m.group(4)) if m.group(4) else now.year
                    end_date  = datetime.date(end_year, end_month, 1)
                start_date = datetime.date(start_year, start_month, 1)
                months = (
                    (end_date.year - start_date.year) * 12
                    + (end_date.month - start_date.month)
                )
                return max(0, months)
            except (ValueError, TypeError):
                pass

        # [C7] Try MM/YYYY range
        m = _MM_YYYY_RANGE_RE.search(date_str)
        if m:
            try:
                start_month = int(m.group(1))
                start_year  = int(m.group(2))
                if m.group(5):  # present/current/now
                    end_date = now
                else:
                    end_month = int(m.group(3))
                    end_year  = int(m.group(4))
                    end_date  = datetime.date(end_year, end_month, 1)
                start_date = datetime.date(start_year, start_month, 1)
                months = (
                    (end_date.year - start_date.year) * 12
                    + (end_date.month - start_date.month)
                )
                return max(0, months)
            except (ValueError, TypeError):
                pass

        # Year-only range
        m = _YEAR_RANGE_RE.search(date_str)
        if m:
            try:
                start_year = int(m.group(1))
                end_token  = m.group(2).lower()
                end_year   = now.year if end_token in ("present", "current", "now") \
                             else int(end_token)
                return max(0, (end_year - start_year) * 12)
            except (ValueError, TypeError):
                pass

        return 0